"""
Extract detailed content from DOCX files:
- Document structure (paragraphs, tables, images in order)
- Paragraph styles, alignment, font properties
- Table structure with merged cells
- Embedded images saved to disk
- Convert to PDF via Word COM
"""
import os
import sys
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import zipfile
from PIL import Image
import io

INPUT_DIR = Path(r"F:\SimlyDent\BM Học kỳ DN T7.2026")
OUTPUT_DIR = INPUT_DIR / "_extracted"
PDF_DIR = INPUT_DIR / "_pdf"

def ensure_dirs():
    OUTPUT_DIR.mkdir(exist_ok=True)
    PDF_DIR.mkdir(exist_ok=True)

def get_alignment_str(alignment):
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }
    return mapping.get(alignment, "left")

def get_font_info(run):
    font = run.font
    info = {}
    if font.name:
        info["name"] = font.name
    if font.size:
        info["size_pt"] = font.size.pt
    info["bold"] = run.bold
    info["italic"] = run.italic
    info["underline"] = run.underline
    if font.color and font.color.rgb:
        info["color"] = str(font.color.rgb)
    return info

def extract_images_from_docx(docx_path, out_folder):
    """Extract all images from the docx zip."""
    images = {}
    out_folder.mkdir(exist_ok=True, parents=True)
    with zipfile.ZipFile(docx_path, 'r') as z:
        for name in z.namelist():
            if name.startswith('word/media/'):
                img_data = z.read(name)
                img_filename = os.path.basename(name)
                img_path = out_folder / img_filename
                with open(img_path, 'wb') as f:
                    f.write(img_data)
                # Get image dimensions
                try:
                    img = Image.open(io.BytesIO(img_data))
                    images[name] = {
                        "saved_as": str(img_path),
                        "filename": img_filename,
                        "width_px": img.width,
                        "height_px": img.height,
                        "format": img.format
                    }
                except:
                    images[name] = {
                        "saved_as": str(img_path),
                        "filename": img_filename,
                    }
    return images

def get_image_rId_map(doc):
    """Map rId to image path in docx."""
    rels = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            rels[rel.rId] = rel.target_ref
    return rels

def extract_paragraph_content(para, rid_map, images_info):
    """Extract detailed paragraph content including inline images."""
    elements = []
    for child in para._element:
        if child.tag == qn('w:r'):
            # Check for inline image
            drawings = child.findall(qn('w:drawing'))
            if drawings:
                for drawing in drawings:
                    blip = drawing.find('.//' + qn('a:blip'))
                    if blip is not None:
                        embed = blip.get(qn('r:embed'))
                        if embed and embed in rid_map:
                            img_ref = 'word/' + rid_map[embed]
                            img_info = images_info.get(img_ref, {})
                            # Get size from drawing
                            extent = drawing.find('.//' + qn('wp:extent'))
                            w_emu = int(extent.get('cx', 0)) if extent is not None else 0
                            h_emu = int(extent.get('cy', 0)) if extent is not None else 0
                            elements.append({
                                "type": "image",
                                "rId": embed,
                                "filename": img_info.get("filename", ""),
                                "saved_as": img_info.get("saved_as", ""),
                                "width_cm": round(w_emu / 914400 * 2.54, 2) if w_emu else None,
                                "height_cm": round(h_emu / 914400 * 2.54, 2) if h_emu else None,
                            })
            else:
                # Text run
                from docx.text.run import Run
                run = None
                for r in para.runs:
                    if r._element is child:
                        run = r
                        break
                if run and run.text:
                    elements.append({
                        "type": "text",
                        "text": run.text,
                        "font": get_font_info(run),
                    })
    return elements

def extract_table(table):
    """Extract table structure."""
    rows_data = []
    for row in table.rows:
        cells_data = []
        for cell in row.cells:
            cell_paras = []
            for p in cell.paragraphs:
                cell_paras.append({
                    "text": p.text,
                    "alignment": get_alignment_str(p.alignment) if p.alignment else "left",
                    "style": p.style.name if p.style else None,
                    "runs": [{"text": r.text, "bold": r.bold, "italic": r.italic, "underline": r.underline} for r in p.runs]
                })
            # Check for merged cells
            tc = cell._tc
            grid_span = tc.find(qn('w:tcPr'))
            colspan = 1
            if grid_span is not None:
                gs = grid_span.find(qn('w:gridSpan'))
                if gs is not None:
                    colspan = int(gs.get(qn('w:val'), 1))
            vmerge = None
            if grid_span is not None:
                vm = grid_span.find(qn('w:vMerge'))
                if vm is not None:
                    vmerge = vm.get(qn('w:val'), 'continue')

            cells_data.append({
                "paragraphs": cell_paras,
                "colspan": colspan,
                "vmerge": vmerge,
            })
        rows_data.append(cells_data)
    return rows_data

def get_section_info(doc):
    """Get page setup info."""
    sections = []
    for sec in doc.sections:
        sections.append({
            "page_width_cm": round(sec.page_width.cm, 2) if sec.page_width else None,
            "page_height_cm": round(sec.page_height.cm, 2) if sec.page_height else None,
            "left_margin_cm": round(sec.left_margin.cm, 2) if sec.left_margin else None,
            "right_margin_cm": round(sec.right_margin.cm, 2) if sec.right_margin else None,
            "top_margin_cm": round(sec.top_margin.cm, 2) if sec.top_margin else None,
            "bottom_margin_cm": round(sec.bottom_margin.cm, 2) if sec.bottom_margin else None,
            "orientation": "landscape" if sec.page_width and sec.page_height and sec.page_width > sec.page_height else "portrait",
        })
    return sections

def extract_body_elements(doc, rid_map, images_info):
    """Extract body elements in document order (paragraphs and tables)."""
    elements = []
    for element in doc.element.body:
        if element.tag == qn('w:p'):
            # Paragraph
            from docx.text.paragraph import Paragraph
            para = Paragraph(element, doc)
            content = extract_paragraph_content(para, rid_map, images_info)
            elements.append({
                "type": "paragraph",
                "text": para.text,
                "style": para.style.name if para.style else None,
                "alignment": get_alignment_str(para.alignment) if para.alignment else None,
                "content": content,
            })
        elif element.tag == qn('w:tbl'):
            # Table
            from docx.table import Table
            table = Table(element, doc)
            elements.append({
                "type": "table",
                "rows": len(table.rows),
                "cols": len(table.columns),
                "data": extract_table(table),
            })
    return elements

def extract_headers_footers(doc):
    """Extract headers and footers."""
    hf = []
    for i, sec in enumerate(doc.sections):
        sec_hf = {"section": i}
        # Header
        if sec.header and not sec.header.is_linked_to_previous:
            sec_hf["header"] = [p.text for p in sec.header.paragraphs]
        # Footer
        if sec.footer and not sec.footer.is_linked_to_previous:
            sec_hf["footer"] = [p.text for p in sec.footer.paragraphs]
        hf.append(sec_hf)
    return hf

def process_docx(docx_path):
    """Process a single DOCX file."""
    name = Path(docx_path).stem
    print(f"\n{'='*60}")
    print(f"Processing: {Path(docx_path).name}")
    print(f"{'='*60}")

    doc = Document(docx_path)

    # Extract images
    img_folder = OUTPUT_DIR / name / "images"
    images_info = extract_images_from_docx(docx_path, img_folder)
    print(f"  Images extracted: {len(images_info)}")

    # Get rId map
    rid_map = get_image_rId_map(doc)

    # Extract sections
    sections = get_section_info(doc)
    print(f"  Sections: {len(sections)}")

    # Extract body
    body_elements = extract_body_elements(doc, rid_map, images_info)
    para_count = sum(1 for e in body_elements if e["type"] == "paragraph")
    table_count = sum(1 for e in body_elements if e["type"] == "table")
    print(f"  Paragraphs: {para_count}, Tables: {table_count}")

    # Extract headers/footers
    hf = extract_headers_footers(doc)

    # Save JSON
    result = {
        "filename": Path(docx_path).name,
        "sections": sections,
        "headers_footers": hf,
        "body": body_elements,
        "images": images_info,
    }

    json_path = OUTPUT_DIR / name / "structure.json"
    json_path.parent.mkdir(exist_ok=True, parents=True)
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print(f"  Saved: {json_path}")

    return result

def convert_to_pdf():
    """Convert all DOCX to PDF using Word COM."""
    print(f"\n{'='*60}")
    print("Converting DOCX to PDF via Microsoft Word...")
    print(f"{'='*60}")
    try:
        import comtypes.client
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        
        docx_files = list(INPUT_DIR.glob("*.docx"))
        for docx_path in docx_files:
            pdf_path = PDF_DIR / (docx_path.stem + ".pdf")
            print(f"  Converting: {docx_path.name} -> {pdf_path.name}")
            try:
                doc = word.Documents.Open(str(docx_path))
                doc.SaveAs(str(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
                doc.Close()
                print(f"    OK")
            except Exception as e:
                print(f"    ERROR: {e}")
        
        word.Quit()
        print("Word closed.")
    except ImportError:
        print("comtypes not available, trying win32com...")
        try:
            import win32com.client
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            
            docx_files = list(INPUT_DIR.glob("*.docx"))
            for docx_path in docx_files:
                pdf_path = PDF_DIR / (docx_path.stem + ".pdf")
                print(f"  Converting: {docx_path.name} -> {pdf_path.name}")
                try:
                    doc = word.Documents.Open(str(docx_path))
                    doc.SaveAs(str(pdf_path), FileFormat=17)
                    doc.Close()
                    print(f"    OK")
                except Exception as e:
                    print(f"    ERROR: {e}")
            
            word.Quit()
            print("Word closed.")
        except ImportError:
            print("ERROR: Neither comtypes nor win32com available.")
            print("Install with: pip install comtypes")
            return False
    return True

def main():
    ensure_dirs()
    
    # Process all DOCX files
    docx_files = sorted(INPUT_DIR.glob("*.docx"))
    print(f"Found {len(docx_files)} DOCX files\n")
    
    all_results = {}
    for docx_path in docx_files:
        try:
            result = process_docx(str(docx_path))
            all_results[docx_path.name] = result
        except Exception as e:
            print(f"  ERROR processing {docx_path.name}: {e}")
            import traceback
            traceback.print_exc()
    
    # Summary
    print(f"\n{'='*60}")
    print("EXTRACTION SUMMARY")
    print(f"{'='*60}")
    for name, result in all_results.items():
        body = result["body"]
        paras = sum(1 for e in body if e["type"] == "paragraph")
        tables = sum(1 for e in body if e["type"] == "table")
        imgs = len(result["images"])
        print(f"  {name}: {paras} paras, {tables} tables, {imgs} images")
    
    # Convert to PDF
    convert_to_pdf()
    
    print("\nDone!")

if __name__ == "__main__":
    main()
